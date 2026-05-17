"""Modulo form_generation: analisi e compilazione di moduli PDF AcroForm.

I template dei moduli sono i PDF originali caricati dall'interfaccia
**Moduli -> Configurazione**. I nomi dei campi compilabili coincidono con i
nomi dei campi AcroForm presenti nel PDF; la mappatura verso i campi database
degli incidenti resta memorizzata nella tabella ``form_field_mapping``.
"""
from __future__ import annotations

import os
import re
import subprocess
import uuid
import html
import base64
import zipfile
from io import BytesIO
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Tuple
from xml.etree import ElementTree as ET

from flask import current_app
from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from pypdf import PdfReader, PdfWriter
from pypdf.generic import NameObject, DictionaryObject, TextStringObject, NumberObject, ArrayObject, IndirectObject

from reportlab.pdfbase.pdfmetrics import stringWidth

from .models import db, Incident, FormFieldMapping, FormTemplateConfig, FormTemplateBinary, Setting

# I campi compilabili dei modelli DOCX sono identificati esclusivamente
# dal costrutto %<testo>%. Il nome del campo è il testo racchiuso tra
# due caratteri % successivi, normalizzato solo per rimuovere spazi ai bordi.
PERCENT_FIELD_RE = re.compile(r"%([^%\r\n]+)%")
FIELD_RE = PERCENT_FIELD_RE
PLACEHOLDER_RE = PERCENT_FIELD_RE

PDF_FONT_RESOURCE_NAMES = {
    'Helvetica': '/Helv',
    'Times-Roman': '/TiRo',
}
PDF_FONT_BASE_NAMES = {
    'Helvetica': '/Helvetica',
    'Times-Roman': '/Times-Roman',
}


def get_template_config(template_name: str) -> FormTemplateConfig:
    cfg = FormTemplateConfig.query.filter_by(template_name=template_name).first()
    if not cfg:
        cfg = FormTemplateConfig(template_name=template_name, font_family='Helvetica', font_size=10)
    cfg.font_family = FormTemplateConfig.normalize_font_family(cfg.font_family)
    cfg.font_size = FormTemplateConfig.normalize_font_size(cfg.font_size)
    return cfg


def save_template_config(template_name: str, font_family: str, font_size) -> FormTemplateConfig:
    safe_name = Path(template_name).stem
    cfg = FormTemplateConfig.query.filter_by(template_name=safe_name).first()
    if not cfg:
        cfg = FormTemplateConfig(template_name=safe_name)
    cfg.font_family = FormTemplateConfig.normalize_font_family(font_family)
    cfg.font_size = FormTemplateConfig.normalize_font_size(font_size)
    return cfg


@dataclass(frozen=True)
class TemplateInfo:
    name: str
    path: Path
    fields: List[str]
    source_docx: Path | None = None
    source_pdf: Path | None = None


def template_dir() -> Path:
    path = Path(current_app.config.get('FORM_TEMPLATE_DIR') or '/data/form_templates')
    path.mkdir(parents=True, exist_ok=True)
    return path


def _fields_from_pdf(path: Path) -> List[str]:
    try:
        fields_dict = PdfReader(str(path)).get_fields() or {}
    except Exception:
        current_app.logger.warning('Impossibile leggere i campi del PDF template %s', path, exc_info=True)
        return []
    fields: List[str] = []
    for name in fields_dict.keys():
        clean = str(name).strip()
        if clean and clean not in fields:
            fields.append(clean)
    return fields


def _restore_template_pdf_from_db(template_name: str) -> Path | None:
    """Ripristina sul file system il PDF template salvato nel DB, se presente."""
    safe = _sanitize_field_name(Path(template_name).stem, 'template')
    row = FormTemplateBinary.query.filter_by(template_name=safe).first()
    if not row or not row.pdf_data:
        return None
    path = template_dir() / f'{safe}.pdf'
    if not path.exists():
        path.write_bytes(row.pdf_data)
        current_app.logger.info('Template PDF %s ripristinato dal database persistente', safe)
    return path


def restore_missing_template_files_from_db() -> int:
    """Ripristina tutti i PDF template mancanti usando la copia binaria nel DB."""
    restored = 0
    for row in FormTemplateBinary.query.order_by(FormTemplateBinary.template_name).all():
        safe = _sanitize_field_name(Path(row.template_name).stem, 'template')
        path = template_dir() / f'{safe}.pdf'
        if row.pdf_data and not path.exists():
            path.write_bytes(row.pdf_data)
            restored += 1
    if restored:
        current_app.logger.info('Ripristinati %s template PDF dal database persistente', restored)
    return restored


def list_templates() -> List[TemplateInfo]:
    restore_missing_template_files_from_db()
    templates: List[TemplateInfo] = []
    seen = set()
    for path in sorted(template_dir().glob('*.pdf')):
        templates.append(TemplateInfo(path.stem, path, _fields_from_pdf(path), None, path))
        seen.add(path.stem)
    for row in FormTemplateBinary.query.order_by(FormTemplateBinary.template_name).all():
        safe = _sanitize_field_name(Path(row.template_name).stem, 'template')
        if safe in seen:
            continue
        path = _restore_template_pdf_from_db(safe)
        if path and path.exists():
            templates.append(TemplateInfo(safe, path, _fields_from_pdf(path), None, path))
            seen.add(safe)
    return templates


def get_template(name: str) -> TemplateInfo:
    safe = _sanitize_field_name(Path(name).stem, 'template')
    path = template_dir() / f"{safe}.pdf"
    if not path.exists():
        _restore_template_pdf_from_db(safe)
    if not path.exists():
        raise FileNotFoundError(f"Template PDF non trovato: {safe}")
    return TemplateInfo(safe, path, _fields_from_pdf(path), None, path)


def save_template_pdf(template_name: str, source_pdf_bytes: bytes) -> Path:
    safe = _sanitize_field_name(Path(template_name).stem, 'template')
    if not source_pdf_bytes:
        raise ValueError('PDF sorgente mancante')
    out = template_dir() / f'{safe}.pdf'
    with open(out, 'wb') as fh:
        fh.write(source_pdf_bytes)

    # Copia persistente nel database: evita la perdita dei template se, dopo
    # un riavvio, il file system montato per i template non contiene più i PDF.
    row = FormTemplateBinary.query.filter_by(template_name=safe).first()
    if not row:
        row = FormTemplateBinary(template_name=safe, filename=f'{safe}.pdf', pdf_data=source_pdf_bytes)
        db.session.add(row)
    else:
        row.filename = f'{safe}.pdf'
        row.pdf_data = source_pdf_bytes
        row.updated_at = datetime.utcnow()
    db.session.commit()
    return out

def _add_unique(target: List[str], value: str | None) -> None:
    value = (value or '').strip()
    if value and value not in target:
        target.append(value)


def _extract_percent_fields(value: str | None) -> List[str]:
    if not value:
        return []
    return [m.group(1).strip() for m in PERCENT_FIELD_RE.finditer(value) if m.group(1).strip()]


def extract_template_fields(path: Path) -> List[str]:
    """Return fillable fields declared in an XML form template.

    Current templates generated from DOCX use <field name="..."/> nodes.
    Hand-written XML templates can also contain literal %FIELD_NAME% tokens in
    text or attributes. Both forms are supported, but the DOCX analyzer creates
    fields only from %...% tokens found in the DOCX source.
    """
    root = ET.parse(path).getroot()
    found: List[str] = []
    for elem in root.iter():
        tag = elem.tag.split('}')[-1].lower()
        if tag == 'field':
            _add_unique(found, elem.attrib.get('name') or elem.attrib.get('id') or elem.attrib.get('field') or (elem.text or ''))
        for name in _extract_percent_fields(elem.text):
            _add_unique(found, name)
        for name in _extract_percent_fields(elem.tail):
            _add_unique(found, name)
        for value in elem.attrib.values():
            for name in _extract_percent_fields(value):
                _add_unique(found, name)
    return found

def available_incident_fields() -> List[Tuple[str, str]]:
    """Campi database disponibili per il mapping drag & drop."""
    return [
        ('name', 'Nome incidente'),
        ('reference', 'Riferimento'),
        ('recipient', 'Destinatario comunicazioni data breach'),
        ('description', 'Descrizione'),
        ('severity', 'Gravità'),
        ('personal_data', 'Dati personali'),
        ('data_subjects_count', 'Numero di interessati'),
        ('data_volume', 'Volume dati'),
        ('start_at', 'Data e ora inizio'),
        ('start_date', 'Data inizio'),
        ('start_time', 'Ora inizio'),
        ('end_at', 'Data e ora fine'),
        ('end_date', 'Data fine'),
        ('end_time', 'Ora fine'),
        ('status', 'Stato'),
        ('creator_name', 'Nome compilatore'),
        ('creator_email', 'Email compilatore'),
        ('categories', 'Categorie'),
        ('category_descriptions', 'Descrizione e causa'),
        ('data_types', 'Dati interessati'),
        ('people', 'Personale coinvolto'),
        ('actions', 'Azioni intraprese'),
        ('awareness_date', 'Data venuta a conoscenza'),
        ('awareness_time', 'Ora venuta a conoscenza'),
        ('documents', 'Documenti allegati'),
        ('created_at', 'Data creazione record'),
        ('security_owner', 'Titolare della sicurezza'),
        ('security_owner_role', 'Ruolo titolare'),
        ('structure', 'Struttura'),
        ('security_responsible', 'Responsabile della sicurezza'),
        ('security_responsible_email', 'Email responsabile'),
        ('security_responsible_phone', 'Telefono responsabile'),
        ('security_responsible_function', 'Funzione responsabile'),
        ('consequences', 'Conseguenze derivate'),
        ('measures_adopted', 'Misure adottate derivate'),
        ('privacy_authority_non_notification_reason', 'Motivi non garante'),
        ('documentation_location', 'Luogo documentazione'),
        ('recommendations', 'Raccomandazioni selezionate'),
    ]



def first_initial_information_action(inc: Incident):
    """Prima azione di tipo 'informazione iniziale' dell'incidente."""
    candidates = []
    for action in inc.actions:
        label = (action.label.value if action.label else '').lower()
        if 'informazione iniziale' in label:
            candidates.append(action)
    return sorted(candidates, key=lambda x: x.when_at or datetime.min)[0] if candidates else None

def _setting_value(key: str, default: str = '') -> str:
    s = Setting.query.get(key)
    return s.value if s and s.value is not None else default

def incident_value(inc: Incident, field_name: str) -> str:
    fmt = lambda d: d.strftime('%Y-%m-%d %H:%M') if d else ''
    if field_name == 'name': return inc.name or ''
    if field_name == 'reference': return inc.reference or ''
    if field_name == 'recipient': return inc.recipient or inc.reference or ''
    if field_name == 'description': return inc.description or ''
    if field_name == 'severity': return inc.severity.value if inc.severity else ''
    if field_name == 'personal_data': return 'Sono presenti dati personali coinvolti.' if inc.personal_data else 'Non risultano dati personali coinvolti.'
    if field_name == 'data_subjects_count': return inc.data_subjects_count or ''
    if field_name == 'data_volume': return inc.data_volume or ''
    if field_name == 'start_at': return fmt(inc.start_at)
    if field_name == 'start_date': return inc.start_date.strftime('%Y-%m-%d') if getattr(inc, 'start_date', None) else (inc.start_at.strftime('%Y-%m-%d') if inc.start_at else '')
    if field_name == 'start_time': return inc.start_time.strftime('%H:%M') if getattr(inc, 'start_time', None) else (inc.start_at.strftime('%H:%M') if inc.start_at else '')
    if field_name == 'end_at': return fmt(inc.end_at)
    if field_name == 'end_date': return inc.end_date.strftime('%Y-%m-%d') if getattr(inc, 'end_date', None) else (inc.end_at.strftime('%Y-%m-%d') if inc.end_at else '')
    if field_name == 'end_time': return inc.end_time.strftime('%H:%M') if getattr(inc, 'end_time', None) else (inc.end_at.strftime('%H:%M') if inc.end_at else '')
    if field_name == 'status': return inc.status or ''
    if field_name == 'creator_name': return inc.creator_name or ''
    if field_name == 'creator_email': return inc.creator_email or ''
    if field_name == 'categories': return ', '.join([c.value for c in inc.categories])
    if field_name == 'category_descriptions': return '\n'.join([c.description or c.value for c in inc.categories])
    if field_name == 'data_types': return ', '.join([d.value for d in inc.data_types])
    if field_name == 'people': return ', '.join([p.name for p in sorted(inc.people, key=lambda x: x.name or '')])
    if field_name == 'actions':
        lines=[]
        for a in sorted(inc.actions, key=lambda x: x.when_at or datetime.min):
            when=fmt(a.when_at)
            label=a.label.value if a.label else ''
            desc=a.description or ''
            lines.append(f"{when} - {label} - {a.person_name or ''}: {desc}".strip())
        return '\n'.join(lines)
    if field_name == 'awareness_date':
        action = first_initial_information_action(inc)
        return action.when_at.strftime('%Y-%m-%d') if action and action.when_at else ''
    if field_name == 'awareness_time':
        action = first_initial_information_action(inc)
        return action.when_at.strftime('%H:%M') if action and action.when_at else ''
    if field_name == 'documents': return ', '.join([d.filename for d in inc.documents])
    if field_name == 'created_at': return fmt(inc.created_at)
    if field_name == 'security_owner': return (Setting.query.get('security_owner_name').value if Setting.query.get('security_owner_name') else '')
    if field_name == 'security_owner_role': return (Setting.query.get('security_owner_role').value if Setting.query.get('security_owner_role') else '')
    if field_name == 'structure': return (Setting.query.get('structure_name').value if Setting.query.get('structure_name') else '')
    if field_name == 'security_responsible': return (Setting.query.get('security_responsible_name').value if Setting.query.get('security_responsible_name') else '')
    if field_name == 'security_responsible_email': return (Setting.query.get('security_responsible_email').value if Setting.query.get('security_responsible_email') else '')
    if field_name == 'security_responsible_phone': return (Setting.query.get('security_responsible_phone').value if Setting.query.get('security_responsible_phone') else '-')
    if field_name == 'security_responsible_function': return (Setting.query.get('security_responsible_function').value if Setting.query.get('security_responsible_function') else '')
    if field_name == 'consequences': return incident_consequences(inc)
    if field_name == 'measures_adopted': return incident_measures(inc)
    if field_name == 'privacy_authority_non_notification_reason': return _setting_value('privacy_authority_non_notification_reason')
    if field_name == 'documentation_location': return _setting_value('documentation_location')
    if field_name == 'recommendations': return '\n'.join([r.text for r in inc.recommendations])
    return ''


def mapping_for_template(template_name: str) -> Dict[str, str]:
    return {m.template_field: m.db_field for m in FormFieldMapping.query.filter_by(template_name=template_name).all()}


def values_for_template(inc: Incident, template_name: str) -> Dict[str, str]:
    mappings = mapping_for_template(template_name)
    return {template_field: incident_value(inc, db_field) for template_field, db_field in mappings.items()}


def incident_field_label_map() -> Dict[str, str]:
    """Etichette leggibili dei campi incidente disponibili per i mapping."""
    return {code: label for code, label in available_incident_fields()}


def missing_required_incident_fields_for_templates(inc: Incident, template_names: Iterable[str]) -> List[Dict[str, object]]:
    """Campi incidente mappati ma privi di valore prima della generazione moduli.

    La generazione dei PDF viene bloccata se uno o più template selezionati
    usano campi database dell'incidente il cui valore risulta vuoto. Il
    controllo è effettuato sui mapping correnti del template, così l'operatore
    riceve un'unica richiesta cumulativa con tutti i campi da completare.
    """
    labels = incident_field_label_map()
    grouped: List[Dict[str, object]] = []
    for template_name in template_names:
        safe_name = _sanitize_field_name(Path(template_name).stem, 'template')
        mappings = mapping_for_template(safe_name)
        missing_by_db_field: Dict[str, Dict[str, object]] = {}
        for template_field, db_field in mappings.items():
            value = incident_value(inc, db_field)
            if str(value or '').strip():
                continue
            item = missing_by_db_field.setdefault(db_field, {
                'db_field': db_field,
                'label': labels.get(db_field, db_field),
                'template_fields': [],
            })
            item['template_fields'].append(template_field)
        if missing_by_db_field:
            grouped.append({
                'template': safe_name,
                'fields': list(missing_by_db_field.values()),
            })
    return grouped


def format_missing_required_incident_fields(missing: List[Dict[str, object]]) -> str:
    """Messaggio sintetico per l'interfaccia utente sui campi mancanti."""
    parts: List[str] = []
    for group in missing:
        field_labels = []
        for field in group.get('fields', []):
            label = str(field.get('label') or field.get('db_field') or '').strip()
            pdf_fields = ', '.join(str(x) for x in field.get('template_fields', []))
            field_labels.append(f"{label} (campo PDF: {pdf_fields})" if pdf_fields else label)
        if field_labels:
            parts.append(f"{group.get('template')}: " + '; '.join(field_labels))
    return 'Prima di generare i moduli completare i campi usati dai template selezionati: ' + ' | '.join(parts)


def _truthy(value: str | None) -> bool:
    return str(value or '').lower() in {'1', 'true', 'yes', 'si', 'sì'}


def _add_run(paragraph, text: str, elem=None):
    run = paragraph.add_run(text or '')
    if elem is not None:
        if _truthy(elem.attrib.get('bold')): run.bold = True
        if _truthy(elem.attrib.get('italic')): run.italic = True
        if _truthy(elem.attrib.get('underline')): run.underline = True
        if elem.attrib.get('size'):
            try: run.font.size = Pt(float(elem.attrib['size']))
            except ValueError: pass
    return run


def _replace_placeholders(text: str, values: Dict[str, str]) -> str:
    return PLACEHOLDER_RE.sub(lambda m: values.get(m.group(1).strip(), ''), text or '')


def _render_inline(paragraph, elem, values: Dict[str, str]):
    if elem.text:
        _add_run(paragraph, _replace_placeholders(elem.text, values), elem)
    for child in list(elem):
        tag = child.tag.split('}')[-1].lower()
        if tag == 'field':
            name = child.attrib.get('name') or child.attrib.get('id') or child.attrib.get('field') or (child.text or '').strip()
            _add_run(paragraph, values.get(name, ''), child)
        elif tag in {'text', 'run', 'span'}:
            _add_run(paragraph, _replace_placeholders(child.text or child.attrib.get('value',''), values), child)
            _render_inline(paragraph, child, values)
        else:
            _render_inline(paragraph, child, values)
        if child.tail:
            _add_run(paragraph, _replace_placeholders(child.tail, values), elem)


def _cell_text(cell_elem, values: Dict[str, str]) -> str:
    chunks=[]
    def walk(e):
        if e.text: chunks.append(_replace_placeholders(e.text, values))
        tag=e.tag.split('}')[-1].lower()
        if tag=='field':
            name=e.attrib.get('name') or e.attrib.get('id') or e.attrib.get('field') or (e.text or '').strip()
            chunks.append(values.get(name,''))
        for c in list(e):
            walk(c)
            if c.tail: chunks.append(_replace_placeholders(c.tail, values))
    walk(cell_elem)
    return ''.join(chunks).strip()



def _iter_document_paragraphs(doc):
    """Yield all paragraphs, including those nested in tables."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    for row2 in nested.rows:
                        for cell2 in row2.cells:
                            for p2 in cell2.paragraphs:
                                yield p2


def _replace_placeholders_in_paragraph_preserving_format(paragraph, values: Dict[str, str]) -> None:
    """Replace %FIELD% tokens while preserving DOCX layout and formatting.

    If a placeholder lives inside a single run, only that run text is replaced,
    keeping its formatting. If Word has split the placeholder across multiple
    runs, the paragraph text is replaced in the first run and the other runs are
    cleared; paragraph/table layout is still preserved and this fallback is
    preferable to losing the entire document structure.
    """
    if not paragraph.runs:
        return
    changed = False
    for run in paragraph.runs:
        original = run.text or ''
        replaced = _replace_placeholders(original, values)
        if replaced != original:
            run.text = replaced
            changed = True
    if changed:
        return
    full = ''.join(run.text or '' for run in paragraph.runs)
    replaced_full = _replace_placeholders(full, values)
    if replaced_full != full:
        paragraph.runs[0].text = replaced_full
        for run in paragraph.runs[1:]:
            run.text = ''


def _generate_docx_from_source_docx(tmpl: TemplateInfo, values: Dict[str, str], output_dir: Path, inc: Incident) -> Path:
    """Generate a DOCX by editing the original DOCX template in place.

    This path preserves the original document formatting much better than
    rebuilding the file from XML: section layout, margins, headers/footers,
    tables, paragraph styles, fonts, images and most run-level formatting are
    kept by python-docx because only placeholder text is replaced.
    """
    doc = DocxDocument(str(tmpl.source_docx))
    for p in _iter_document_paragraphs(doc):
        _replace_placeholders_in_paragraph_preserving_format(p, values)
    # Headers and footers may contain placeholders in official letter templates.
    for section in doc.sections:
        for container in (section.header, section.footer, section.first_page_header, section.first_page_footer,
                          section.even_page_header, section.even_page_footer):
            for p in container.paragraphs:
                _replace_placeholders_in_paragraph_preserving_format(p, values)
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            _replace_placeholders_in_paragraph_preserving_format(p, values)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{tmpl.name}-{inc.id}-{uuid.uuid4().hex[:8]}.docx"
    out = output_dir / filename
    doc.save(out)
    return out

def generate_docx_from_xml(inc: Incident, template_name: str, output_dir: Path) -> Path:
    tmpl = get_template(template_name)
    values = values_for_template(inc, template_name)
    if tmpl.source_docx and tmpl.source_docx.exists():
        return _generate_docx_from_source_docx(tmpl, values, output_dir, inc)
    root = ET.parse(tmpl.path).getroot()
    doc = DocxDocument()
    title = root.attrib.get('title') or template_name
    doc.add_heading(title, level=1)
    for elem in list(root):
        tag = elem.tag.split('}')[-1].lower()
        if tag in {'title', 'heading'}:
            doc.add_heading(_cell_text(elem, values) or elem.attrib.get('text',''), level=int(elem.attrib.get('level','2')))
        elif tag in {'paragraph', 'p'}:
            p = doc.add_paragraph(style=elem.attrib.get('style') if elem.attrib.get('style') else None)
            _render_inline(p, elem, values)
        elif tag == 'table':
            rows = [r for r in list(elem) if r.tag.split('}')[-1].lower() == 'row']
            if not rows: continue
            max_cols = max(len([c for c in list(r) if c.tag.split('}')[-1].lower() == 'cell']) for r in rows)
            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.style = elem.attrib.get('style', 'Table Grid')
            for ri, row in enumerate(rows):
                cells = [c for c in list(row) if c.tag.split('}')[-1].lower() == 'cell']
                for ci, cell_elem in enumerate(cells):
                    table.cell(ri, ci).text = _cell_text(cell_elem, values)
        else:
            p = doc.add_paragraph()
            _render_inline(p, elem, values)
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{template_name}-{inc.id}-{uuid.uuid4().hex[:8]}.docx"
    out = output_dir / filename
    doc.save(out)
    return out


def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    try:
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', str(output_dir), str(docx_path)
        ], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=120)
        pdf = output_dir / (docx_path.stem + '.pdf')
        if pdf.exists() and pdf.stat().st_size > 0:
            return pdf
    except Exception as exc:
        current_app.logger.warning('Conversione DOCX->PDF con LibreOffice fallita: %s', exc)
    # Fallback: PDF informativo se LibreOffice non è disponibile. Il container
    # ufficiale installa LibreOffice, quindi questo ramo serve solo a evitare crash.
    pdf = output_dir / (docx_path.stem + '.pdf')
    styles=getSampleStyleSheet()
    story=[Paragraph('Conversione PDF non disponibile', styles['Title']), Spacer(1, 12), Paragraph(f'Documento DOCX generato: {docx_path.name}', styles['BodyText'])]
    SimpleDocTemplate(str(pdf), pagesize=A4).build(story)
    return pdf


# ---------------------------------------------------------------------------
# PDF form template extraction / preview / generation
# ---------------------------------------------------------------------------

def analyze_pdf_template(pdf_path: str | Path, title: str | None = None) -> Tuple[List[str], List[Dict[str, str]], str, List[Dict[str, str]]]:
    """Analyze a fillable PDF form and return AcroForm field names.

    The uploaded PDF itself is saved as template. No intermediate XML template
    is created: the field names shown here are the real PDF field names used
    later during generation.
    """
    pdf_path = Path(pdf_path)
    reader = PdfReader(str(pdf_path))
    fields_dict = reader.get_fields() or {}
    fields: List[str] = []
    field_meta: List[Dict[str, str]] = []
    for name, meta in fields_dict.items():
        clean = str(name).strip()
        if not clean or clean in fields:
            continue
        fields.append(clean)
        ft = ''
        try:
            ft = str(meta.get('/FT', '') or '') if hasattr(meta, 'get') else ''
        except Exception:
            ft = ''
        field_meta.append({'name': clean, 'source': 'campo modulo PDF', 'sample': ft})
    elements = [{'type': 'PDF', 'text': f'{len(fields)} campi modulo rilevati'}]
    visual_html = _render_pdf_fields_preview_html(fields, field_meta)
    return fields, elements, visual_html, field_meta


def _render_pdf_fields_preview_html(fields: List[str], field_meta: List[Dict[str, str]]) -> str:
    out = ['<div class="docx-visual-page"><h3>Anteprima campi modulo PDF</h3>']
    if not fields:
        out.append('<p class="procedure-warning"><strong>Nessun campo modulo PDF rilevato.</strong><span>Usare un PDF con campi AcroForm compilabili.</span></p>')
    else:
        out.append('<table class="docx-preview-table"><tr><th>Nome campo PDF</th><th>Tipo</th></tr>')
        for item in field_meta:
            out.append(f'<tr><td><span class="docx-field-marker">{html.escape(item.get("name", ""))}</span></td><td>{html.escape(item.get("sample", ""))}</td></tr>')
        out.append('</table>')
    out.append('</div>')
    return ''.join(out)


PDF_TEXT_MULTILINE_FLAG = 1 << 12
PDF_MIN_FIELD_FONT_SIZE = 8.0
PDF_MAX_FIELD_FONT_SIZE = 16.0
PDF_FIELD_PADDING = 4.0


def _pdf_text_width(text: str, font_size: float, font_family: str = 'Helvetica') -> float:
    try:
        return stringWidth(text or '', FormTemplateConfig.normalize_font_family(font_family), font_size)
    except Exception:
        # Conservative fallback when ReportLab metrics are unavailable.
        return len(text or '') * font_size * 0.55


def _wrap_text_for_width(text: str, usable_width: float, font_size: float, font_family: str = 'Helvetica') -> List[str]:
    """Wrap text using the full available field width.

    Existing newlines are preserved as paragraph breaks. Very long words are
    split at character level so that each generated line fits inside the field
    width as much as possible. The function returns at least one line.
    """
    usable_width = max(usable_width, font_size * 2)
    result: List[str] = []
    for paragraph in str(text or '').replace('\r\n', '\n').replace('\r', '\n').split('\n'):
        words = paragraph.split()
        if not words:
            result.append('')
            continue
        line = ''
        for word in words:
            candidate = word if not line else f'{line} {word}'
            if _pdf_text_width(candidate, font_size, font_family) <= usable_width:
                line = candidate
                continue
            if line:
                result.append(line)
                line = ''
            if _pdf_text_width(word, font_size, font_family) <= usable_width:
                line = word
                continue
            chunk = ''
            for ch in word:
                candidate_chunk = chunk + ch
                if chunk and _pdf_text_width(candidate_chunk, font_size, font_family) > usable_width:
                    result.append(chunk)
                    chunk = ch
                else:
                    chunk = candidate_chunk
            line = chunk
        result.append(line)
    return result or ['']


def _field_rect_size(obj) -> Tuple[float, float]:
    try:
        rect = obj.get('/Rect') if hasattr(obj, 'get') else None
        if not rect:
            return 0.0, 0.0
        vals = [float(x) for x in rect]
        return abs(vals[2] - vals[0]), abs(vals[3] - vals[1])
    except Exception:
        return 0.0, 0.0


def _choose_font_and_text(value: str, width: float, height: float, font_size: int, font_family: str) -> Tuple[float, str, bool]:
    """Return configured font size and wrapped text using the available width."""
    value = str(value or '')
    font_family = FormTemplateConfig.normalize_font_family(font_family)
    font_size = FormTemplateConfig.normalize_font_size(font_size)
    if not value:
        return float(font_size), value, False

    usable_width = max(width - (PDF_FIELD_PADDING * 2), font_size * 4) if width else 240.0
    # The configured size is respected; wrapping uses the full field width.
    lines = _wrap_text_for_width(value, usable_width, float(font_size), font_family)
    multiline = len(lines) > 1 or '\n' in value
    return float(font_size), '\n'.join(lines), multiline


def _ensure_pdf_font_resource(acro_obj, font_family: str) -> str:
    """Ensure the selected standard PDF font is available in AcroForm resources."""
    family = FormTemplateConfig.normalize_font_family(font_family)
    resource_name = PDF_FONT_RESOURCE_NAMES.get(family, '/Helv')
    base_name = PDF_FONT_BASE_NAMES.get(family, '/Helvetica')
    try:
        dr = acro_obj.get('/DR')
        if not dr:
            dr = DictionaryObject()
            acro_obj.update({NameObject('/DR'): dr})
        dr_obj = dr.get_object() if hasattr(dr, 'get_object') else dr
        fonts = dr_obj.get('/Font')
        if not fonts:
            fonts = DictionaryObject()
            dr_obj.update({NameObject('/Font'): fonts})
        fonts_obj = fonts.get_object() if hasattr(fonts, 'get_object') else fonts
        key = NameObject(resource_name)
        if key not in fonts_obj:
            fonts_obj.update({key: DictionaryObject({
                NameObject('/Type'): NameObject('/Font'),
                NameObject('/Subtype'): NameObject('/Type1'),
                NameObject('/BaseFont'): NameObject(base_name),
            })})
    except Exception:
        pass
    return resource_name

def _set_pdf_field_default_appearance(obj, font_size: float, font_resource_name: str = '/Helv') -> None:
    """Set default appearance using configured font and size."""
    try:
        da = obj.get('/DA') if hasattr(obj, 'get') else None
        if da:
            da_text = str(da)
            da_text = re.sub(r'/(\S+)\s+[-+]?\d+(?:\.\d+)?\s+Tf', f'{font_resource_name} {font_size:g} Tf', da_text)
            if da_text == str(da):
                da_text = f'{font_resource_name} {font_size:g} Tf 0 g'
            obj.update({NameObject('/DA'): TextStringObject(da_text)})
        else:
            obj.update({NameObject('/DA'): TextStringObject(f'{font_resource_name} {font_size:g} Tf 0 g')})
    except Exception:
        pass


def _enable_multiline_if_needed(obj, multiline: bool) -> None:
    if not multiline:
        return
    try:
        flags = int(obj.get('/Ff', 0) or 0) if hasattr(obj, 'get') else 0
        obj.update({NameObject('/Ff'): NumberObject(flags | PDF_TEXT_MULTILINE_FLAG)})
    except Exception:
        pass


def _is_text_widget(obj) -> bool:
    try:
        ft = obj.get('/FT') if hasattr(obj, 'get') else None
        return str(ft) in {'/Tx', 'Tx'} or ft is None
    except Exception:
        return True


def _prepare_pdf_text_fields_for_wrapping(writer: PdfWriter, values: Dict[str, str], font_family: str, font_size: int) -> Dict[str, str]:
    """Prepare text widgets for readable multiline filling.

    Pypdf writes the form values but relies on the viewer to regenerate visual
    appearances. Setting /DA and /Ff on both field dictionaries and widget
    annotations makes common viewers render wrapped, readable text using the
    configured font family and size selected for the template.
    """
    adjusted_values = dict(values)
    field_specs: Dict[str, Tuple[float, str, bool]] = {}
    font_family = FormTemplateConfig.normalize_font_family(font_family)
    font_size = FormTemplateConfig.normalize_font_size(font_size)
    font_resource_name = '/Helv'

    try:
        acro = writer._root_object.get('/AcroForm')
        if acro:
            acro_obj = acro.get_object() if hasattr(acro, 'get_object') else acro
            if '/DR' not in acro_obj:
                acro_obj.update({NameObject('/DR'): DictionaryObject()})
            font_resource_name = _ensure_pdf_font_resource(acro_obj, font_family)
    except Exception:
        pass

    def prepare_obj(obj, inherited_name: str | None = None):
        try:
            name = str(obj.get('/T') or inherited_name or '').strip()
            if not name or name not in adjusted_values or not _is_text_widget(obj):
                return
            width, height = _field_rect_size(obj)
            chosen_size, wrapped_value, multiline = _choose_font_and_text(adjusted_values.get(name, ''), width, height, font_size, font_family)
            prev = field_specs.get(name)
            # Prefer specs from actual widgets with dimensions; otherwise keep first.
            if prev is None or (width and height):
                field_specs[name] = (chosen_size, wrapped_value, multiline)
                adjusted_values[name] = wrapped_value
            _set_pdf_field_default_appearance(obj, chosen_size, font_resource_name)
            _enable_multiline_if_needed(obj, multiline)
        except Exception:
            current_app.logger.debug('Preparazione campo PDF non completata', exc_info=True)

    try:
        acro = writer._root_object.get('/AcroForm')
        if acro:
            acro_obj = acro.get_object() if hasattr(acro, 'get_object') else acro
            for field_ref in acro_obj.get('/Fields', []):
                field = field_ref.get_object() if hasattr(field_ref, 'get_object') else field_ref
                prepare_obj(field)
    except Exception:
        pass

    for page in writer.pages:
        try:
            for annot_ref in page.get('/Annots') or []:
                annot = annot_ref.get_object() if hasattr(annot_ref, 'get_object') else annot_ref
                if annot.get('/Subtype') == '/Widget':
                    parent_name = None
                    parent = annot.get('/Parent')
                    if parent:
                        try:
                            parent_obj = parent.get_object() if hasattr(parent, 'get_object') else parent
                            parent_name = str(parent_obj.get('/T') or '').strip()
                        except Exception:
                            parent_name = None
                    prepare_obj(annot, parent_name)
                    name = str(annot.get('/T') or parent_name or '').strip()
                    if name in field_specs:
                        font_size, wrapped_value, multiline = field_specs[name]
                        _set_pdf_field_default_appearance(annot, font_size, font_resource_name)
                        _enable_multiline_if_needed(annot, multiline)
                        adjusted_values[name] = wrapped_value
        except Exception:
            pass
    return adjusted_values

def _pdf_object_key(obj):
    """Stable key for a PDF object, also when it is indirect."""
    try:
        if isinstance(obj, IndirectObject):
            return (int(obj.idnum), int(obj.generation))
    except Exception:
        pass
    try:
        indirect = getattr(obj, 'indirect_reference', None)
        if indirect is not None:
            return (int(indirect.idnum), int(indirect.generation))
    except Exception:
        pass
    return ('mem', id(obj))


def _field_name_map(reader: PdfReader) -> Dict[object, str]:
    """Map AcroForm field/widget objects to their fully qualified field name.

    Complex PDFs often store the visible widget annotation as a child of one or
    more parent field dictionaries. The page annotation may therefore only carry
    a partial /T value, or no /T at all. Using the fully qualified AcroForm path
    keeps mappings reliable for nested fields such as ``section.field`` and for
    fields with several widgets.
    """
    names: Dict[object, str] = {}
    try:
        root = reader.trailer.get('/Root') or {}
        acro = root.get('/AcroForm') if hasattr(root, 'get') else None
        acro_obj = acro.get_object() if hasattr(acro, 'get_object') else acro
        fields = acro_obj.get('/Fields', []) if hasattr(acro_obj, 'get') else []
    except Exception:
        fields = []

    def visit(ref, prefix: str = '') -> None:
        try:
            obj = ref.get_object() if hasattr(ref, 'get_object') else ref
            part = str(obj.get('/T') or '').strip() if hasattr(obj, 'get') else ''
            full = f'{prefix}.{part}' if prefix and part else (part or prefix)
            if full:
                names[_pdf_object_key(ref)] = full
                names[_pdf_object_key(obj)] = full
            for kid_ref in obj.get('/Kids', []) or []:
                visit(kid_ref, full)
        except Exception:
            current_app.logger.debug('Impossibile risolvere nome campo PDF annidato', exc_info=True)

    for field_ref in fields or []:
        visit(field_ref, '')
    return names


def _annot_field_name(annot, name_map: Dict[object, str] | None = None) -> str:
    """Return the effective, preferably fully qualified, field name."""
    name_map = name_map or {}
    try:
        mapped = name_map.get(_pdf_object_key(annot))
        if mapped:
            return mapped
    except Exception:
        pass
    parts = []
    current = annot
    seen = set()
    while current is not None:
        try:
            obj = current.get_object() if hasattr(current, 'get_object') else current
            key = _pdf_object_key(obj)
            if key in seen:
                break
            seen.add(key)
            mapped = name_map.get(key)
            if mapped:
                return mapped
            part = str(obj.get('/T') or '').strip() if hasattr(obj, 'get') else ''
            if part:
                parts.append(part)
            current = obj.get('/Parent') if hasattr(obj, 'get') else None
        except Exception:
            break
    if parts:
        return '.'.join(reversed(parts))
    return ''


def _annot_field_type(annot) -> str:
    """Return the effective AcroForm field type for a widget annotation."""
    try:
        current = annot
        seen = set()
        while current is not None:
            obj = current.get_object() if hasattr(current, 'get_object') else current
            key = _pdf_object_key(obj)
            if key in seen:
                break
            seen.add(key)
            ft = obj.get('/FT') if hasattr(obj, 'get') else None
            if ft:
                return str(ft)
            current = obj.get('/Parent') if hasattr(obj, 'get') else None
    except Exception:
        pass
    return ''


def _annotation_is_visible_for_flattening(annot) -> bool:
    try:
        flags = int(annot.get('/F', 0) or 0)
        # PDF annotation flags: Invisible=1, Hidden=2, NoView=32.
        return not (flags & 1 or flags & 2 or flags & 32)
    except Exception:
        return True

def _lookup_field_value(values: Dict[str, str], field_name: str):
    """Return the configured value for a widget name, accepting safe aliases.

    Some PDFs expose full hierarchical names (for example
    ``section.field``) while older mappings or specific editors may store only
    the terminal widget name.  The exact name remains preferred, then the
    terminal part is accepted as a fallback to avoid silently skipping visible
    fields during final-print generation.
    """
    if field_name in values:
        return values.get(field_name)
    terminal = field_name.split('.')[-1] if field_name else ''
    if terminal and terminal in values:
        return values.get(terminal)
    lowered = field_name.lower()
    for key, value in values.items():
        if str(key).lower() == lowered or str(key).lower() == terminal.lower():
            return value
    return None


def _effective_rect(annot):
    try:
        current = annot
        seen = set()
        while current is not None:
            obj = current.get_object() if hasattr(current, 'get_object') else current
            key = _pdf_object_key(obj)
            if key in seen:
                break
            seen.add(key)
            rect = obj.get('/Rect') if hasattr(obj, 'get') else None
            if rect:
                return rect
            current = obj.get('/Parent') if hasattr(obj, 'get') else None
    except Exception:
        pass
    return None


def _draw_static_pdf_field(c, value: str, rect, font_family: str, font_size: int, field_type: str = '') -> None:
    """Draw a field value directly on the PDF page canvas.

    This simulates a final print: values are painted as page content rather
    than left inside interactive AcroForm widgets.  The renderer deliberately
    accepts very small fields: it reduces the drawing font when necessary and
    vertically centers single-line values instead of dropping them.
    """
    value = str(value or '')
    if not value:
        return
    try:
        x1, y1, x2, y2 = [float(v) for v in rect]
    except Exception:
        return
    left, right = min(x1, x2), max(x1, x2)
    bottom, top = min(y1, y2), max(y1, y2)
    width, height = right - left, top - bottom
    if width <= 0 or height <= 0:
        return

    family = FormTemplateConfig.normalize_font_family(font_family)
    configured_size = float(FormTemplateConfig.normalize_font_size(font_size))
    # Keep the configured size where possible, but never let a short field
    # disappear merely because its rectangle is smaller than the selected font.
    size = max(4.0, min(configured_size, max(height - (PDF_FIELD_PADDING * 2), height * 0.85)))
    c.setFont(family, size)

    if str(field_type) in {'/Btn', 'Btn'}:
        mark = 'X' if value.lower() in {'1', 'true', 'yes', 'si', 'sì', 'on'} else value
        c.drawCentredString(left + width / 2, bottom + max((height - size) / 2, 0.5), mark[:4])
        return

    _, wrapped, _ = _choose_font_and_text(value, width, height, int(round(size)), family)
    lines = wrapped.split('\n') if wrapped else []
    if not lines:
        return
    leading = max(size * 1.15, size + 1)
    x = left + min(PDF_FIELD_PADDING, max(width * 0.08, 0.5))
    min_y = bottom + 0.5

    if len(lines) == 1:
        # ReportLab's y coordinate is the text baseline; center the baseline
        # inside the rectangle so single-line date/time fields remain visible.
        y = bottom + max((height - size) / 2, 0.5)
        c.drawString(x, y, lines[0])
        return

    y = top - min(PDF_FIELD_PADDING, max(height * 0.12, 0.5)) - size
    if y < min_y:
        y = bottom + max(height - size, 0.5)
    for line in lines:
        if y < min_y:
            break
        c.drawString(x, y, line)
        y -= leading


def _overlay_for_page(page, widgets, values: Dict[str, str], font_family: str, font_size: int, name_map: Dict[object, str] | None = None) -> PdfReader | None:
    """Create a one-page PDF overlay with static field values for a page."""
    if not widgets:
        return None
    try:
        box = page.cropbox if getattr(page, 'cropbox', None) else page.mediabox
        width = float(box.width)
        height = float(box.height)
        offset_x = float(box.left)
        offset_y = float(box.bottom)
    except Exception:
        try:
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            offset_x = 0.0
            offset_y = 0.0
        except Exception:
            return None
    packet = BytesIO()
    c = canvas.Canvas(packet, pagesize=(width, height))
    if offset_x or offset_y:
        c.translate(-offset_x, -offset_y)
    drawn = set()
    for annot in widgets:
        annot_obj = annot.get_object() if hasattr(annot, 'get_object') else annot
        if not _annotation_is_visible_for_flattening(annot_obj):
            continue
        name = _annot_field_name(annot, name_map) or _annot_field_name(annot_obj, name_map)
        if not name:
            continue
        value = _lookup_field_value(values, name)
        if value is None:
            continue
        rect = _effective_rect(annot_obj)
        if not rect:
            continue
        draw_key = (_pdf_object_key(annot), str(name), tuple(str(x) for x in rect))
        if draw_key in drawn:
            continue
        drawn.add(draw_key)
        _draw_static_pdf_field(c, value, rect, font_family, font_size, _annot_field_type(annot_obj))
    c.save()
    packet.seek(0)
    return PdfReader(packet)




def _collect_acroform_widgets_by_page(reader: PdfReader) -> Dict[object, List[object]]:
    """Collect widget annotations from the AcroForm tree grouped by page.

    Relying only on ``page['/Annots']`` misses fields produced by some PDF
    editors where the widget is present in the AcroForm ``/Kids`` tree but is
    not consistently repeated in the page annotation array.  Grouping by the
    widget ``/P`` page reference lets the final-print pass draw those fields as
    well.
    """
    by_page: Dict[object, List[object]] = {}
    page_keys = {}
    try:
        for page in reader.pages:
            page_keys[_pdf_object_key(page)] = page
            ref = getattr(page, 'indirect_reference', None)
            if ref is not None:
                page_keys[_pdf_object_key(ref)] = page
    except Exception:
        return by_page

    try:
        acro = (reader.trailer.get('/Root') or {}).get('/AcroForm')
        acro_obj = acro.get_object() if hasattr(acro, 'get_object') else acro
        fields = acro_obj.get('/Fields', []) if hasattr(acro_obj, 'get') else []
    except Exception:
        fields = []

    def visit(ref):
        try:
            obj = ref.get_object() if hasattr(ref, 'get_object') else ref
            subtype = obj.get('/Subtype') if hasattr(obj, 'get') else None
            if subtype == '/Widget' or obj.get('/Rect'):
                page_ref = obj.get('/P') if hasattr(obj, 'get') else None
                page = page_keys.get(_pdf_object_key(page_ref)) if page_ref is not None else None
                if page is not None:
                    by_page.setdefault(_pdf_object_key(page), []).append(ref)
            for kid_ref in obj.get('/Kids', []) or []:
                visit(kid_ref)
        except Exception:
            current_app.logger.debug('Widget AcroForm non collezionato', exc_info=True)

    for field_ref in fields or []:
        visit(field_ref)
    return by_page


def _append_unique_widget(widgets: List[object], widget) -> None:
    try:
        key = _pdf_object_key(widget)
        obj = widget.get_object() if hasattr(widget, 'get_object') else widget
        obj_key = _pdf_object_key(obj)
        for existing in widgets:
            existing_obj = existing.get_object() if hasattr(existing, 'get_object') else existing
            if _pdf_object_key(existing) == key or _pdf_object_key(existing_obj) == obj_key:
                return
        widgets.append(widget)
    except Exception:
        widgets.append(widget)

def generate_pdf_from_template(inc: Incident, template_name: str, output_dir: Path) -> Path:
    """Generate a filled, flattened PDF from the original uploaded form.

    The output is a new final-print PDF: mapped values are drawn as static page
    content and AcroForm widget annotations are removed, so the generated file
    no longer contains fillable fields.
    """
    tmpl = get_template(template_name)
    if not tmpl.source_pdf or not tmpl.source_pdf.exists():
        raise FileNotFoundError('Il template selezionato non contiene un PDF sorgente compilabile')
    values = values_for_template(inc, template_name)
    cfg = get_template_config(tmpl.name)
    reader = PdfReader(str(tmpl.source_pdf))
    name_map = _field_name_map(reader)
    acro_widgets_by_page = _collect_acroform_widgets_by_page(reader)
    writer = PdfWriter()

    for page in reader.pages:
        widgets = []
        other_annots = ArrayObject()
        try:
            for annot_ref in page.get('/Annots') or []:
                annot = annot_ref.get_object() if hasattr(annot_ref, 'get_object') else annot_ref
                if annot.get('/Subtype') == '/Widget':
                    # Keep the indirect reference when possible: it is the most
                    # reliable key for inherited field names in nested AcroForms.
                    _append_unique_widget(widgets, annot_ref)
                else:
                    other_annots.append(annot_ref)
        except Exception:
            widgets = []
            other_annots = page.get('/Annots') or ArrayObject()

        for extra_widget in acro_widgets_by_page.get(_pdf_object_key(page), []):
            _append_unique_widget(widgets, extra_widget)

        overlay_reader = _overlay_for_page(page, widgets, values, cfg.font_family, cfg.font_size, name_map)
        if overlay_reader and overlay_reader.pages:
            try:
                page.merge_page(overlay_reader.pages[0])
            except Exception:
                current_app.logger.debug('Flatten overlay PDF non applicato per %s', template_name, exc_info=True)

        # Remove AcroForm widget annotations from the generated page while
        # preserving unrelated annotations such as links or comments.
        try:
            if other_annots:
                page[NameObject('/Annots')] = other_annots
            elif '/Annots' in page:
                del page['/Annots']
        except Exception:
            pass
        writer.add_page(page)

    # Do not copy /AcroForm into the output: the file must be non-interactive.
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{tmpl.name}-{inc.id}-{uuid.uuid4().hex[:8]}.pdf"
    out = output_dir / filename
    with open(out, 'wb') as fh:
        writer.write(fh)
    return out

# ---------------------------------------------------------------------------
# DOCX -> XML template extraction / preview
# ---------------------------------------------------------------------------
from xml.dom import minidom

# Durante l'analisi DOCX non vengono più interpretati campi grigi, puntini,
# sottolineature o vecchi nomi input_data. Sono compilabili solo i token %...%.
FIELD_LIKE_RE = PERCENT_FIELD_RE


def _xml_string(obj) -> str:
    try:
        return obj.xml
    except Exception:
        try:
            return obj._element.xml
        except Exception:
            return ''


def _is_gray_or_highlighted_run(run) -> bool:
    """Return True when a run looks like a fillable grey field.

    Word documents may mark fields either using text highlight or XML shading
    (w:shd/w:highlight). The exact color can vary across templates; for the
    purpose of form generation every highlighted/shaded run is presented to the
    user as a candidate fillable field and can be renamed before saving.
    """
    try:
        if run.font.highlight_color is not None:
            return True
    except Exception:
        pass
    x = _xml_string(run._r).lower()
    if 'w:highlight' in x and 'w:val="none"' not in x:
        return True
    if 'w:shd' in x and 'w:fill="auto"' not in x and 'w:fill="ffffff"' not in x:
        return True
    return False


def _is_shaded_cell(cell) -> bool:
    x = _xml_string(cell._tc).lower()
    return 'w:shd' in x and 'w:fill="auto"' not in x and 'w:fill="ffffff"' not in x


def _sanitize_field_name(value: str, fallback: str) -> str:
    value = re.sub(r'[^A-Za-z0-9_\-]+', '_', (value or '').strip()).strip('_-')
    return value or fallback

def _format_attrs_from_run(run) -> Dict[str, str]:
    attrs: Dict[str, str] = {}
    if run.bold:
        attrs['bold'] = 'true'
    if run.italic:
        attrs['italic'] = 'true'
    if run.underline:
        attrs['underline'] = 'true'
    try:
        if run.font.size:
            attrs['size'] = str(round(run.font.size.pt, 1))
    except Exception:
        pass
    return attrs


def _pretty_xml(elem: ET.Element) -> str:
    rough = ET.tostring(elem, encoding='utf-8')
    return minidom.parseString(rough).toprettyxml(indent='  ', encoding='utf-8').decode('utf-8')


def _field_name_from_token(token: str, counter: int) -> str:
    match = PERCENT_FIELD_RE.fullmatch(token or '')
    if match:
        return match.group(1).strip() or f'field_{counter:03d}'
    return f'field_{counter:03d}'


def _append_run_as_template(parent: ET.Element, run, state: Dict[str, int], fields: List[str], field_meta: List[Dict[str, str]] | None = None) -> None:
    """Append a Word run to XML, turning only %...% tokens into fields."""
    text_value = run.text or ''
    if not text_value:
        return
    attrs = _format_attrs_from_run(run)
    pos = 0
    for match in FIELD_LIKE_RE.finditer(text_value):
        before = text_value[pos:match.start()]
        if before:
            ET.SubElement(parent, 'text', attrs).text = before
        token = match.group(0)
        state['counter'] += 1
        name = _field_name_from_token(token, state['counter'])
        if name not in fields:
            fields.append(name)
            if field_meta is not None:
                field_meta.append({'name': name, 'source': 'placeholder %...%', 'sample': token[:120]})
        ET.SubElement(parent, 'field', {'name': name, **attrs})
        pos = match.end()
    after = text_value[pos:]
    if after:
        ET.SubElement(parent, 'text', attrs).text = after

def _paragraph_to_template_xml(root: ET.Element, paragraph, state: Dict[str, int], fields: List[str], preview: List[Dict[str, str]], field_meta: List[Dict[str, str]] | None = None) -> None:
    style_name = getattr(getattr(paragraph, 'style', None), 'name', '') or ''
    tag = 'heading' if style_name.lower().startswith('heading') or style_name.lower().startswith('titolo') else 'paragraph'
    attrs = {'style': style_name} if style_name else {}
    elem = ET.SubElement(root, tag, attrs)
    text_preview = paragraph.text or ''
    if text_preview.strip():
        preview.append({'type': 'Paragrafo', 'text': text_preview[:180]})
    for run in paragraph.runs:
        _append_run_as_template(elem, run, state, fields, field_meta)
    if not list(elem) and text_preview:
        ET.SubElement(elem, 'text').text = text_preview


def _table_to_template_xml(root: ET.Element, table, state: Dict[str, int], fields: List[str], preview: List[Dict[str, str]], field_meta: List[Dict[str, str]] | None = None) -> None:
    table_elem = ET.SubElement(root, 'table', {'style': getattr(getattr(table, 'style', None), 'name', '') or 'Table Grid'})
    preview.append({'type': 'Tabella', 'text': f'{len(table.rows)} righe × {len(table.columns)} colonne'})
    for row in table.rows:
        row_elem = ET.SubElement(table_elem, 'row')
        for cell in row.cells:
            cell_elem = ET.SubElement(row_elem, 'cell')
            shaded_cell = _is_shaded_cell(cell)
            cell_had_content = False
            for p in cell.paragraphs:
                para_elem = ET.SubElement(cell_elem, 'paragraph')
                for run in p.runs:
                    before_children = len(list(para_elem))
                    _append_run_as_template(para_elem, run, state, fields, field_meta)
                    if len(list(para_elem)) > before_children or (run.text or '').strip():
                        cell_had_content = True
                if not list(para_elem) and p.text:
                    ET.SubElement(para_elem, 'text').text = p.text
                    cell_had_content = True
            # Le celle grigie senza token %...% non sono considerate campi compilabili.



def _render_xml_to_preview_html(xml_content: str) -> str:
    """Render generated XML as a lightweight graphical preview for the UI."""
    root = ET.fromstring(xml_content.encode('utf-8'))
    out = ['<div class="docx-visual-page">']

    def inline(e) -> str:
        chunks = []
        if e.text:
            chunks.append(html.escape(e.text))
        for c in list(e):
            tag = c.tag.split('}')[-1].lower()
            if tag == 'field':
                name = html.escape(c.attrib.get('name', 'input_data'))
                chunks.append(f'<span class="docx-field-marker">{name}</span>')
            else:
                chunks.append(inline(c))
            if c.tail:
                chunks.append(html.escape(c.tail))
        return ''.join(chunks)

    for elem in list(root):
        tag = elem.tag.split('}')[-1].lower()
        if tag in {'heading', 'title'}:
            out.append(f'<h3>{inline(elem)}</h3>')
        elif tag in {'paragraph', 'p'}:
            out.append(f'<p>{inline(elem) or "&nbsp;"}</p>')
        elif tag == 'table':
            out.append('<table class="docx-preview-table">')
            for row in [r for r in list(elem) if r.tag.split('}')[-1].lower() == 'row']:
                out.append('<tr>')
                for cell in [c for c in list(row) if c.tag.split('}')[-1].lower() == 'cell']:
                    out.append(f'<td>{inline(cell) or "&nbsp;"}</td>')
                out.append('</tr>')
            out.append('</table>')
        else:
            out.append(f'<p>{inline(elem)}</p>')
    out.append('</div>')
    return ''.join(out)


def rename_fields_in_xml(xml_content: str, renames: Dict[str, str]) -> str:
    """Rename field identifiers in an XML template before saving."""
    root = ET.fromstring(xml_content.encode('utf-8'))
    clean = {old: _sanitize_field_name(new, old) for old, new in renames.items() if old and new}
    for elem in root.iter():
        if elem.tag.split('}')[-1].lower() == 'field':
            name = elem.attrib.get('name')
            if name in clean:
                elem.set('name', clean[name])
        if elem.text:
            for old, new in clean.items():
                elem.text = elem.text.replace(old, new)
        if elem.tail:
            for old, new in clean.items():
                elem.tail = elem.tail.replace(old, new)
        for key, value in list(elem.attrib.items()):
            if key == 'pdf_name':
                continue
            for old, new in clean.items():
                value = value.replace(old, new)
            elem.set(key, value)
    return _pretty_xml(root)

def extract_docx_to_template_xml(docx_path: str | Path, title: str | None = None) -> Tuple[str, List[str], List[Dict[str, str]], str, List[Dict[str, str]]]:
    """Create an internal XML form template from an uploaded DOCX.

    The conversion preserves a practical subset of DOCX structure used by the
    generator: paragraphs, headings, tables and run formatting attributes
    (bold/italic/underline/font size). Fillable fields are detected when the
    source document contains tokens del tipo ``%<testo>%``. Puntini, campi grigi o sottolineature non sono più considerati campi compilabili.
    """
    doc = DocxDocument(str(docx_path))
    root = ET.Element('formTemplate', {'title': title or Path(docx_path).stem, 'source_docx': f"{_sanitize_field_name(title or Path(docx_path).stem, 'template')}.docx", 'format_preservation': 'source_docx'})
    fields: List[str] = []
    preview: List[Dict[str, str]] = []
    field_meta: List[Dict[str, str]] = []
    state = {'counter': 0}

    # Preserve the original document order for block items.
    body = doc.element.body
    paragraph_by_element = {p._p: p for p in doc.paragraphs}
    table_by_element = {t._tbl: t for t in doc.tables}
    for child in body.iterchildren():
        if child in paragraph_by_element:
            _paragraph_to_template_xml(root, paragraph_by_element[child], state, fields, preview, field_meta)
        elif child in table_by_element:
            _table_to_template_xml(root, table_by_element[child], state, fields, preview, field_meta)

    # Some DOCX files expose tables/paragraphs in a way that does not match the
    # direct body iterator above; fallback avoids creating empty templates.
    if not list(root):
        for p in doc.paragraphs:
            _paragraph_to_template_xml(root, p, state, fields, preview, field_meta)
        for t in doc.tables:
            _table_to_template_xml(root, t, state, fields, preview, field_meta)

    xml_content = _pretty_xml(root)
    return xml_content, fields, preview, _render_xml_to_preview_html(xml_content), field_meta


def safe_template_filename(template_name: str) -> str:
    safe = secure = re.sub(r'[^A-Za-z0-9_.\-]+', '_', (template_name or '').strip()).strip('._-')
    if not safe:
        raise ValueError('Nome template non valido')
    if not safe.endswith('.xml'):
        safe += '.xml'
    return safe


def save_template_xml(template_name: str, xml_content: str, source_docx_bytes: bytes | None = None, source_pdf_bytes: bytes | None = None) -> Path:
    filename = safe_template_filename(template_name)
    # Validate XML and ensure at least a formTemplate root before saving.
    root = ET.fromstring(xml_content.encode('utf-8'))
    if root.tag.split('}')[-1] != 'formTemplate':
        raise ValueError('Il template XML deve avere root <formTemplate>')
    path = template_dir() / filename
    if source_docx_bytes:
        docx_name = Path(filename).with_suffix('.docx').name
        root.set('source_docx', docx_name)
        root.set('format_preservation', 'source_docx')
        (template_dir() / docx_name).write_bytes(source_docx_bytes)
        xml_content = _pretty_xml(root)
    if source_pdf_bytes:
        pdf_name = Path(filename).with_suffix('.pdf').name
        root.set('source_pdf', pdf_name)
        root.set('template_type', 'pdf_acroform')
        (template_dir() / pdf_name).write_bytes(source_pdf_bytes)
        xml_content = _pretty_xml(root)
    with open(path, 'w', encoding='utf-8') as fh:
        fh.write(xml_content)
    return path


def incident_consequences(inc: Incident) -> str:
    explicit=[a.consequence_text.strip() for a in sorted(inc.actions, key=lambda x: x.when_at or datetime.min) if getattr(a, 'consequence_text', None) and a.consequence_text.strip()]
    if explicit:
        return '\n'.join(explicit)
    cats=[(c.value or '').lower() for c in inc.categories]
    data=[(d.value or '').lower() for d in inc.data_types]
    out=[]
    if any('credential' in c or 'credenzial' in c for c in cats) or any('password' in d for d in data):
        out.append('Possibile compromissione di credenziali, accessi non autorizzati e necessità di rotazione password.')
    if any('phishing' in c for c in cats):
        out.append('Possibile esposizione a messaggi fraudolenti, furto di informazioni o propagazione dell’attacco.')
    if any('spam' in c for c in cats):
        out.append('Possibile ricezione o invio di comunicazioni indesiderate e impatto sulla reputazione dei servizi.')
    if inc.personal_data or any('dati personali' in d for d in data):
        out.append('Possibile coinvolgimento di dati personali con impatti sui diritti e le libertà degli interessati.')
    return '\n'.join(out) if out else 'Conseguenze da valutare sulla base dell’analisi dell’incidente.'

def incident_measures(inc: Incident) -> str:
    lines=[]
    for a in sorted([x for x in inc.actions if getattr(x, 'exportable', True)], key=lambda x: x.when_at or datetime.min):
        label=(a.label.description or a.label.value) if a.label else 'azione'
        desc=a.description or ''
        when=a.when_at.strftime('%Y-%m-%d %H:%M') if a.when_at else ''
        action_text = f'{label}: {desc}'.strip() if desc else label
        lines.append(f'{action_text} - {when}'.strip(' -'))
    return '\n'.join(lines) if lines else 'Nessuna misura registrata.'
